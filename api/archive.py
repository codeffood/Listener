import json
import hashlib
import shutil
from pathlib import Path
from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from fastapi.responses import StreamingResponse
import io

router = APIRouter(prefix="/api/archive", tags=["archive"])

UPLOAD_DIR   = Path("data/uploads")
ARCHIVE_DIR  = Path("data/archive")
NAS_CACHE_DIR = Path("data/cache/nas")
LIBRARY_PATH = Path("data/library.json")


# ── helpers ───────────────────────────────────────────────────────────────────

def _load_library() -> list:
    if LIBRARY_PATH.exists():
        with open(LIBRARY_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return []


def _save_library(entries: list):
    LIBRARY_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(LIBRARY_PATH, "w", encoding="utf-8") as f:
        json.dump(entries, f, ensure_ascii=False, indent=2)


def _nas_cache_path(nas_idx: int, remote_path: str) -> Path:
    key = f"{nas_idx}:{remote_path}"
    h = hashlib.md5(key.encode()).hexdigest()
    return NAS_CACHE_DIR / f"{h}.json"


def _local_cache_path(name: str) -> Path:
    return UPLOAD_DIR / (Path(name).stem + ".json")


def _archived_audio_path(folder: str, name: str) -> Path:
    return ARCHIVE_DIR / folder / name


def _archived_cache_path(folder: str, name: str) -> Path:
    stem = Path(name).stem
    return ARCHIVE_DIR / folder / f"{stem}.json"


# ── list folders ──────────────────────────────────────────────────────────────

@router.get("/folders")
def list_folders():
    ARCHIVE_DIR.mkdir(parents=True, exist_ok=True)
    return [d.name for d in sorted(ARCHIVE_DIR.iterdir()) if d.is_dir()]


# ── list archived entries ─────────────────────────────────────────────────────

@router.get("/entries")
def list_archived():
    entries = _load_library()
    result = []
    for e in entries:
        if not e.get("archived"):
            continue
        ec = dict(e)
        if e["type"] == "local":
            audio = _archived_audio_path(e["archive_folder"], e["name"])
            cache = _archived_cache_path(e["archive_folder"], e["name"])
            ec["exists"]      = audio.exists()
            ec["transcribed"] = cache.exists()
            if audio.exists():
                ec["size"] = audio.stat().st_size
        else:
            ec["transcribed"] = _nas_cache_path(e["nas_idx"], e["path"]).exists()
        result.append(ec)
    return result


# ── archive ───────────────────────────────────────────────────────────────────

class ArchiveRequest(BaseModel):
    folder: str          # destination folder name under data/archive/
    name: str   = ""     # local file name
    path: str   = ""     # NAS remote path
    nas_idx: int = -1


@router.post("/archive")
def archive_entry(req: ArchiveRequest):
    folder = req.folder.strip()
    if not folder:
        raise HTTPException(400, "folder 不能为空")

    dest_dir = ARCHIVE_DIR / folder
    dest_dir.mkdir(parents=True, exist_ok=True)

    entries = _load_library()

    if req.name:
        # ── local file ──
        src_audio = UPLOAD_DIR / req.name
        src_cache = _local_cache_path(req.name)
        if not src_audio.exists():
            raise HTTPException(404, f"文件不存在: {req.name}")

        dst_audio = dest_dir / req.name
        dst_cache = dest_dir / (Path(req.name).stem + ".json")

        shutil.move(str(src_audio), dst_audio)
        if src_cache.exists():
            shutil.move(str(src_cache), dst_cache)

        for e in entries:
            if e["type"] == "local" and e["name"] == req.name:
                e["archived"]       = True
                e["archive_folder"] = folder
                break

    elif req.path and req.nas_idx >= 0:
        # ── NAS entry — move transcription cache only ──
        src_cache = _nas_cache_path(req.nas_idx, req.path)
        if src_cache.exists():
            dst_cache = dest_dir / src_cache.name
            shutil.move(str(src_cache), dst_cache)

        for e in entries:
            if e["type"] == "nas" and e["path"] == req.path:
                e["archived"]            = True
                e["archive_folder"]      = folder
                e["archive_cache_name"]  = src_cache.name
                break
    else:
        raise HTTPException(400, "需要提供 name（本地文件）或 path+nas_idx（NAS）")

    _save_library(entries)
    return {"status": "archived", "folder": folder}


# ── restore ───────────────────────────────────────────────────────────────────

class RestoreRequest(BaseModel):
    name: str   = ""
    path: str   = ""
    nas_idx: int = -1


@router.post("/restore")
def restore_entry(req: RestoreRequest):
    entries = _load_library()

    if req.name:
        entry = next((e for e in entries if e["type"] == "local" and e["name"] == req.name), None)
        if not entry:
            raise HTTPException(404, "库中找不到该条目")

        folder    = entry.get("archive_folder", "")
        src_audio = _archived_audio_path(folder, req.name)
        src_cache = _archived_cache_path(folder, req.name)
        dst_audio = UPLOAD_DIR / req.name
        dst_cache = _local_cache_path(req.name)

        if src_audio.exists():
            UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
            shutil.move(str(src_audio), dst_audio)
        if src_cache.exists():
            shutil.move(str(src_cache), dst_cache)

        entry.pop("archived", None)
        entry.pop("archive_folder", None)

    elif req.path and req.nas_idx >= 0:
        entry = next((e for e in entries if e["type"] == "nas" and e["path"] == req.path), None)
        if not entry:
            raise HTTPException(404, "库中找不到该条目")

        folder          = entry.get("archive_folder", "")
        cache_name      = entry.get("archive_cache_name", "")
        src_cache       = ARCHIVE_DIR / folder / cache_name
        dst_cache       = NAS_CACHE_DIR / cache_name

        if src_cache.exists():
            NAS_CACHE_DIR.mkdir(parents=True, exist_ok=True)
            shutil.move(str(src_cache), dst_cache)

        entry.pop("archived", None)
        entry.pop("archive_folder", None)
        entry.pop("archive_cache_name", None)
    else:
        raise HTTPException(400, "需要提供 name 或 path+nas_idx")

    _save_library(entries)
    return {"status": "restored"}


# ── export segments ───────────────────────────────────────────────────────────

class ExportRequest(BaseModel):
    filename: str
    segments: list   # [{id, start, end, text}, ...]
    fmt: str = "txt" # "txt" | "docx"


def _fmt_time(s: float) -> str:
    m = int(s) // 60
    sec = s - m * 60
    return f"{m:02d}:{sec:05.2f}"


@router.post("/export")
def export_segments(req: ExportRequest):
    title = Path(req.filename).stem

    if req.fmt == "docx":
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise HTTPException(500, "python-docx 未安装，请在 requirements.txt 中添加 python-docx 并重建镜像")

        doc = Document()

        # Title
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # spacer

        for seg in req.segments:
            ts = doc.add_paragraph()
            ts.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_ts = ts.add_run(f"[{_fmt_time(seg['start'])} → {_fmt_time(seg['end'])}]")
            run_ts.font.size = Pt(9)
            run_ts.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            body = doc.add_paragraph()
            run_body = body.add_run(seg["text"].strip())
            run_body.font.size = Pt(12)
            body.paragraph_format.space_after = Pt(8)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        safe = title.replace(" ", "_")
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe}.docx"'},
        )

    else:
        # plain text
        lines = [title, "=" * len(title), ""]
        for seg in req.segments:
            lines.append(f"[{_fmt_time(seg['start'])} → {_fmt_time(seg['end'])}]")
            lines.append(seg["text"].strip())
            lines.append("")
        content = "\n".join(lines)
        safe = title.replace(" ", "_")
        return StreamingResponse(
            iter([content.encode("utf-8")]),
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{safe}.txt"'},
        )
